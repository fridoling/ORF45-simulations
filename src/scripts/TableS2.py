from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
import pickle
import re

document = Document("../../res/TableS2_template.docx")
table = document.tables[0]

pars = [cell.text for cell in table.columns[-1].cells]
colnames = [cell.text for cell in table.rows[0].cells]

## updating SPR parameters
ind_col_SPR = colnames.index('SPR fit^')

with open("../../data/SPR/table_pars_SPR.pickle", 'rb') as f:
    pars_SPR = pickle.load(f)
cells = table.columns[ind_col_SPR].cells
for i, par in zip(range(len(pars)), pars):
    if par in pars_SPR.keys():
        bold = cells[i].paragraphs[0].runs[0].font.bold
        cells[i].text = "{:.2g}".format(pars_SPR[par])
        cells[i].paragraphs[0].runs[0].font.bold = bold
        cells[i].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

## updating invitro parameters
ind_col_invitro = colnames.index('in vitro fit^')

with open("../../data/invitro/table_pars_invitro.pickle", 'rb') as f:
    pars_invitro = pickle.load(f)
cells = table.columns[ind_col_invitro].cells
for i, par in zip(range(len(pars)), pars):
    if par in pars_invitro.keys():
        bold = cells[i].paragraphs[0].runs[0].font.bold
        cells[i].text = "{:.2g}".format(pars_invitro[par])
        cells[i].paragraphs[0].runs[0].font.bold = bold
        cells[i].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER


document.save("../../res/TableS2.docx")
